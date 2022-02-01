import os
import glob
import pandas as pd

from numpy import full
##documentation says it only works on python 3.4 - I'm using 3.8
import docx

def print_all_tables(document: docx.Document):
    for ti, table in enumerate(document.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                print(ti, ri, ci, cell.text)

def print_all_paragraphs(document: docx.Document):
    for pi, paragraph in enumerate(document.paragraphs):
        print(pi, paragraph.text)

class Person:
    def __init__(self, filename):
        self.first = None
        parts = filename.split('_')
        if parts[0] !='' and parts[0]!=' ':
            self.first = parts[0]
            self.last = parts[1]
        else:
            self.first = parts[1]
            self.last = parts[2]
    def extract_school_name(self, filename: str) -> str:
        filename = filename.replace(self.first, "")
        filename = filename.replace(self.last, "")
        filename = filename.replace("Transcript", "")
        filename = filename.replace("transcript", "")
        parts = filename.split("_")
        return " ".join(parts)


def get_ind_transcript_totals(document: docx.Document):
    search_str = "**Transcript Totals**"
    for pi, paragraph in enumerate(document.paragraphs):
        if search_str in paragraph.text:
            return pi+1
    return None

def extract_cum_credit_hours(line: str) -> float:
    label = "GPA Hours:"
    assert label in line, f"expected {label} in line. passed in the wrong line."
    tmp = line.replace(label, "")
    parts = tmp.split()
    return float(parts[2])
def extract_cum_gpa(line):
    label = "AGPA:"
    assert label in line, f"expected {label} in line. passed in the wrong line."
    tmp = line.replace(label, "")
    parts = tmp.split()
    return float(parts[0])
    
def load_from_doc(dir_input: str = os.getcwd()) -> pd.DataFrame:
    """looks for the word documents in the dir input.
    finds the gpa information from each and returns a dataframe"""
    filepath_list = []
    df = pd.DataFrame(columns=["filename","first_name", "last_name" ,"school_name", "cum_credit_hours", "cum_gpa"])
    for ext in [".doc", ".docx"]:
        pat = os.path.join(dir_input, "*"+ext)
        matching = glob.glob(pat)
        filepath_list.extend(matching)
    for filepath in filepath_list:
        filename = filepath.split(os.path.sep)[-1]
        print("processing", filename)
        full_name = Person(filename)
        school_name = full_name.extract_school_name(filename)
        doc = docx.Document(filepath)
        #print_all_paragraphs(doc)
        ind = get_ind_transcript_totals(doc)
        if ind:
            cum_credit_hours = extract_cum_credit_hours(doc.paragraphs[ind].text)
            cum_gpa = extract_cum_gpa(doc.paragraphs[ind+1].text)
        else: 
            cum_credit_hours =-1
            cum_gpa = -1
        row = {"filename":filename,"first_name":full_name.first, "last_name": full_name.last,"school_name": school_name, "cum_credit_hours":cum_credit_hours, "cum_gpa":cum_gpa}
        rowdf = pd.DataFrame(row, index=[0])
        df = pd.concat([df, rowdf], ignore_index=True)
        #df = df.append(row, ignore_index=True)
    return df

def load_from_pdf(dir_input = os.getcwd()):
    """looks for the pdf documents in the dir input.
    finds the gpa information from each and returns a dataframe"""
    filepath_list = []
    df = pd.DataFrame(columns=["filename","first_name", "last_name" ,"school_name", "cum_credit_hours", "cum_gpa"])
    for ext in [".doc", ".docx"]:
        pat = os.path.join(dir_input, "*"+ext)
        matching = glob.glob(pat)
        filepath_list.extend(matching)
    for filepath in filepath_list:
        ##intention is to use PyPDF2
        pass
if __name__=="__main__":
    df =   load_from_doc()  
    print(df)
